import streamlit as st
import pandas as pd
import re
import tempfile
import os
import base64
from io import BytesIO

# Fuzzy matching
from rapidfuzz import process, fuzz

# Geração de DOCX
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Extração de texto do PDF
from PyPDF2 import PdfReader
import fitz

###############################################################################
# FALLBACK PARA st.session_state (EVITA KeyError)
###############################################################################
_fallback_state = {
    "df_consolidado": None,
    "df_filtrado_descontos": None,
    "df_gloss": None,
    "df_incluido": None,
    "nome_cliente": None,
    "matricula": None,
    "nome_servidor": None,
    # Opcionalmente, valor_recebido (B)
    "valor_recebido": "0"
}

def get_state_value(key):
    """Retorna valor do st.session_state ou do fallback, caso a chave não exista."""
    try:
        return st.session_state[key]
    except:
        return _fallback_state.get(key, None)

def set_state_value(key, value):
    """Seta o valor tanto em st.session_state quanto em fallback, para evitar KeyError."""
    try:
        st.session_state[key] = value
    except:
        _fallback_state[key] = value

###############################################################################
# FUNÇÃO PARA REMOVER PREFIXOS/LIXO
###############################################################################
def remover_prefixos_indesejados(texto: str) -> str:
    """
    Remove palavras específicas (ex.: 'CPE', 'AM', 'EST' etc.) que eventualmente
    aparecem nos nomes de servidores, e retorna a string limpa.
    """
    if not texto:
        return "N/D"
    palavras_descartar = [
        "CPE", "AM", "EST", "APOSENTADO", "DEPLIGPORT",
        "ANO", "REFERÊNCIA", "AP", "SIAPE", "COADI"
    ]
    tokens = texto.split()
    tokens_filtrados = []
    for t in tokens:
        t_limp = re.sub(r"[^\wÀ-ÖØ-öø-ÿ]", "", t, flags=re.IGNORECASE)
        if t_limp.upper() in palavras_descartar:
            continue
        tokens_filtrados.append(t)
    retorno = " ".join(tokens_filtrados)
    retorno = re.sub(r"\s+", " ", retorno).strip()
    return retorno

###############################################################################
# EXTRAIR NOME DO CLIENTE (UTILIZANDO PyMuPDF)
###############################################################################
def extrair_nome_cliente(pdf_path):
    """
    Tenta capturar o nome do servidor (e CPF) nas linhas
    após 'NOME DO SERVIDOR', via PyMuPDF (fitz).
    """
    try:
        doc = fitz.open(pdf_path)
        texto_completo = ""
        for page in doc:
            texto_completo += page.get_text("text") + "\n"
        doc.close()

        pattern = re.compile(
            r"NOME\s+DO\s+SERVIDOR.*?(?:\n.*?){0,6}([A-Za-zÀ-ÖØ-öø-ÿ\s]+)\s+(\d{3}\.\d{3}\.\d{3}-\d{2})",
            flags=re.IGNORECASE | re.DOTALL
        )
        match = pattern.search(texto_completo)
        if match:
            nome_bruto = " ".join(match.group(1).split())
            nome_puro = remover_prefixos_indesejados(nome_bruto)
            return nome_puro
        return "N/D"
    except:
        return "N/D"

###############################################################################
# EXTRAIR NOME E MATRÍCULA (UTILIZANDO PyPDF2)
###############################################################################
def extrair_nome_e_matricula(pdf_path):
    """
    Lê a primeira página via PyPDF2 e procura linhas com 'NOME' e 'MATRÍCULA-SEQ-DIG'.
    Retorna 'N/D' se não encontradas.
    """
    nome = "N/D"
    matricula = "N/D"
    try:
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            if len(reader.pages) > 0:
                text = reader.pages[0].extract_text() or ""
                lines = text.split('\n')
                for i, linha in enumerate(lines):
                    if "NOME" in linha.upper() and i+1 < len(lines):
                        valor_nome = lines[i+1].strip()
                        match_nome = re.match(r"([^\d]+)", valor_nome)
                        if match_nome:
                            nome = match_nome.group(1).strip()
                    if "MATRÍCULA-SEQ-DIG" in linha.upper() and i+1 < len(lines):
                        valor_matr = lines[i+1].strip()
                        matr_match = re.search(r"(\d{3}\.\d{3}-\d\s*[A-Z]*)", valor_matr)
                        if matr_match:
                            matricula = matr_match.group(1).strip()
    except:
        pass
    return nome or "N/D", matricula or "N/D"

###############################################################################
# FUNÇÕES DE SUPORTE
###############################################################################
def sanitizar_para_arquivo(texto: str) -> str:
    """
    Substitui espaços por underscore e remove caracteres especiais para criação de nomes de arquivo.
    """
    texto = texto.strip().replace(" ", "_")
    return re.sub(r"[^\w\-_\.]", "", texto, flags=re.UNICODE)

def get_image_base64(file_path):
    """
    Retorna a string base64 de uma imagem, útil para exibir no Streamlit.
    """
    if not os.path.exists(file_path):
        return ""
    with open(file_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

def carregar_glossario(path):
    """
    Lê um arquivo txt (uma rubrica por linha) e retorna como lista de strings.
    """
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().splitlines()
    except Exception as e:
        st.error(f"Erro ao carregar glossário: {e}")
        return []

###############################################################################
# EXTRAIR CELULAS DE INTERESSE (ANO REFERÊNCIA)
###############################################################################
def extrair_celulas_interesse(pdf_path, termo_referencia="ANO REFERÊNCIA"):
    """
    Utiliza Camelot para identificar as células que contêm o texto 'ANO REFERÊNCIA' (ou outro termo).
    Retorna DataFrame com as páginas e o conteúdo.
    """
    import camelot
    try:
        tables = camelot.read_pdf(pdf_path, pages="all", flavor="lattice")
        if not tables.n:
            return None
        dados = []
        for table in tables:
            pagina_atual = int(table.page)
            df_table = table.df.copy()
            for row_idx in range(df_table.shape[0]):
                for col_idx in range(df_table.shape[1]):
                    conteudo_celula = str(df_table.iat[row_idx, col_idx])
                    if termo_referencia.upper() in conteudo_celula.upper():
                        dados.append({
                            "PÁGINA": pagina_atual,
                            "CONTEÚDO": conteudo_celula
                        })
        df_resultado = pd.DataFrame(dados, columns=["PÁGINA", "CONTEÚDO"])
        return df_resultado if not df_resultado.empty else None
    except Exception as e:
        st.error(f"Erro ao extrair células de interesse: {e}")
        return None

def extrair_ultimos_quatro_digitos(texto):
    """
    Retorna os últimos 4 dígitos encontrados em 'texto'.
    """
    numeros = re.findall(r"\b(\d{4})\b", texto)
    return numeros[-1] if numeros else ""

def classificar_registros_ffill(df):
    """
    Preenche valores vazios na coluna 'TIPO' usando forward fill.
    """
    if "TIPO" in df.columns:
        df["TIPO"] = df["TIPO"].replace("", None).ffill()
    return df

###############################################################################
# EXTRAIR TABELAS – DATAFRAME CONSOLIDADO (TODAS AS COLUNAS + ANO)
###############################################################################
def extrair_tabelas(pdf_path, anos_referencia):
    """
    Extrai tabelas de cada página usando Camelot (flavor 'lattice'),
    reorganiza as colunas, e identifica colunas de acordo com página ímpar/par.
    """
    import camelot
    try:
        tables = camelot.read_pdf(pdf_path, pages="all", flavor="lattice")
        if not tables.n:
            st.error("Nenhuma tabela detectada no PDF.")
            return None

        colunas_finais = [
            "PÁGINA", "TIPO", "DISCRIMINAÇÃO",
            "JAN", "FEV", "MAR", "ABR", "MAI", "JUN",
            "JUL", "AGO", "SET", "OUT", "NOV", "DEZ",
            "ANO"
        ]
        df_final = pd.DataFrame(columns=colunas_finais)

        for table in tables:
            df_tab = table.df.copy()
            pagina_atual = int(table.page)

            # Localiza indices de início e fim (com base em 'TIPO' e 'TOTAL BRUTO')
            start_idx_list = df_tab.index[df_tab.apply(
                lambda row: any("TIPO" in str(cell).upper() for cell in row), axis=1
            )].tolist()
            end_idx_list = df_tab.index[df_tab.apply(
                lambda row: any("TOTAL BRUTO" in str(cell).upper() for cell in row), axis=1
            )].tolist()

            if not start_idx_list or not end_idx_list:
                continue
            start_idx = start_idx_list[0]
            end_idx = end_idx_list[0]
            if end_idx <= start_idx:
                continue

            # Recorta o trecho da tabela
            df_slice = df_tab.iloc[start_idx:end_idx].copy()
            df_slice.columns = df_slice.iloc[0].values  # Usar a primeira linha como cabeçalho
            df_slice = df_slice[1:]  # Remove a linha de cabeçalho duplicada
            df_slice.reset_index(drop=True, inplace=True)

            # Ajusta colunas conforme página ímpar ou par
            if pagina_atual % 2 != 0:
                # Páginas ímpares – colunas de JAN a JUN
                colunas_impar = ["TIPO", "DISCRIMINAÇÃO", "JAN", "FEV", "MAR", "ABR", "MAI", "JUN"]
                map_rename = {}
                for c in df_slice.columns:
                    c_up = c.upper().strip()
                    if "TIPO" in c_up:
                        map_rename[c] = "TIPO"
                    elif "DISCRIMIN" in c_up:
                        map_rename[c] = "DISCRIMINAÇÃO"
                    elif "JAN" in c_up:
                        map_rename[c] = "JAN"
                    elif "FEV" in c_up:
                        map_rename[c] = "FEV"
                    elif "MAR" in c_up:
                        map_rename[c] = "MAR"
                    elif "ABR" in c_up:
                        map_rename[c] = "ABR"
                    elif "MAI" in c_up:
                        map_rename[c] = "MAI"
                    elif "JUN" in c_up:
                        map_rename[c] = "JUN"
                df_slice.rename(columns=map_rename, inplace=True)
                for col in colunas_impar:
                    if col not in df_slice.columns:
                        df_slice[col] = None
                for mes in ["JUL", "AGO", "SET", "OUT", "NOV", "DEZ"]:
                    df_slice[mes] = None
            else:
                # Páginas pares – colunas de JUL a DEZ
                colunas_par = ["TIPO", "DISCRIMINAÇÃO", "JUL", "AGO", "SET", "OUT", "NOV", "DEZ"]
                map_rename = {}
                for c in df_slice.columns:
                    c_up = c.upper().strip()
                    if "TIPO" in c_up:
                        map_rename[c] = "TIPO"
                    elif "DISCRIMIN" in c_up:
                        map_rename[c] = "DISCRIMINAÇÃO"
                    elif "JUL" in c_up:
                        map_rename[c] = "JUL"
                    elif "AGO" in c_up:
                        map_rename[c] = "AGO"
                    elif "SET" in c_up:
                        map_rename[c] = "SET"
                    elif "OUT" in c_up:
                        map_rename[c] = "OUT"
                    elif "NOV" in c_up:
                        map_rename[c] = "NOV"
                    elif "DEZ" in c_up:
                        map_rename[c] = "DEZ"
                df_slice.rename(columns=map_rename, inplace=True)
                for col in colunas_par:
                    if col not in df_slice.columns:
                        df_slice[col] = None
                for mes in ["JAN", "FEV", "MAR", "ABR", "MAI", "JUN"]:
                    df_slice[mes] = None

            df_slice["PÁGINA"] = pagina_atual
            df_slice["ANO"] = anos_referencia.get(pagina_atual, "")
            df_slice = df_slice[
                ["PÁGINA", "TIPO", "DISCRIMINAÇÃO",
                 "JAN", "FEV", "MAR", "ABR", "MAI", "JUN",
                 "JUL", "AGO", "SET", "OUT", "NOV", "DEZ", "ANO"]
            ]
            df_final = pd.concat([df_final, df_slice], ignore_index=True)

        return df_final if not df_final.empty else None
    except Exception as e:
        st.error(f"Erro ao extrair tabelas: {e}")
        return None

###############################################################################
# SALVAR DATAFRAME CONSOLIDADO EM PDF – INCLUINDO CABEÇALHO "Extrato Financeiro Único"
###############################################################################
def salvar_em_pdf(df, nome_pdf):
    """
    Gera um PDF com o cabeçalho 'Extrato Financeiro Único' (usando FPDF).
    Cada página do PDF corresponde à página do DF. Ajusta colunas conforme ímpar/par.
    """
    from fpdf import FPDF

    class PDFCustom(FPDF):
        def header(self):
            self.set_font("Arial", "B", 16)
            # Cabeçalho com o título solicitado:
            self.cell(0, 10, "Extrato Financeiro Único", border=False, ln=True, align='C')
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Página {self.page_no()}', border=False, ln=False, align='C')

    pdf = PDFCustom(orientation='L', format='A4')
    pdf.add_page()
    pdf.set_font("Arial", size=10)

    paginas_unicas = sorted(df["PÁGINA"].unique())
    for pagina in paginas_unicas:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, f"Página {pagina}", border=False, ln=True, align='C')
        pdf.ln(5)
        df_pag = df[df["PÁGINA"] == pagina].copy()

        # Remoção das colunas não pertinentes para cada página ímpar/par
        if pagina % 2 != 0:
            for col_drop in ["JUL", "AGO", "SET", "OUT", "NOV", "DEZ"]:
                if col_drop in df_pag.columns:
                    df_pag.drop(columns=col_drop, inplace=True)
        else:
            for col_drop in ["JAN", "FEV", "MAR", "ABR", "MAI", "JUN"]:
                if col_drop in df_pag.columns:
                    df_pag.drop(columns=col_drop, inplace=True)

        colunas_cabecalho = list(df_pag.columns)
        larguras = []
        for col in colunas_cabecalho:
            if col.upper() == "DISCRIMINAÇÃO":
                larguras.append(70)
            elif col.upper() == "PÁGINA":
                larguras.append(15)
            elif col.upper() in ["JAN", "FEV", "MAR", "ABR", "MAI", "JUN",
                                 "JUL", "AGO", "SET", "OUT", "NOV", "DEZ"]:
                larguras.append(20)
            elif col.upper() == "TIPO":
                larguras.append(30)
            elif col.upper() == "ANO":
                larguras.append(15)
            else:
                larguras.append(20)

        pdf.set_font("Arial", size=10)
        # Cabeçalho da tabela
        for i, col in enumerate(colunas_cabecalho):
            pdf.cell(larguras[i], 10, col, border=1, align='C')
        pdf.ln()

        # Linhas
        for _, row in df_pag.iterrows():
            for i, col in enumerate(colunas_cabecalho):
                valor = str(row[col]) if pd.notnull(row[col]) else ""
                pdf.cell(larguras[i], 10, valor, border=1, align='C')
            pdf.ln()
        pdf.ln(5)

    pdf.output(nome_pdf)

###############################################################################
# FUNÇÕES PARA ANÁLISE DE DESCONTOS E GERAÇÃO DE RELATÓRIOS
###############################################################################
def cruzar_descontos_com_rubricas(df_descontos, glossary, threshold=85):
    """
    Faz um fuzzy matching entre as discriminações de df_descontos e o glossary,
    retornando somente as linhas que atingirem a pontuação (threshold) definida.
    """
    if df_descontos.empty or not glossary:
        return pd.DataFrame()
    unique_desc = df_descontos["DISCRIMINAÇÃO"].unique()
    mapping = {}
    for desc in unique_desc:
        result = process.extractOne(desc, glossary, scorer=fuzz.ratio)
        score = result[1] if result else 0
        mapping[desc] = (score >= threshold)
    return df_descontos[df_descontos["DISCRIMINAÇÃO"].map(mapping)]

def formatar_valor_brl(us_string: str) -> str:
    """
    Converte string no formato US (ex.: '123,456.78' ou '1234.56') para BR '123.456,78'.
    Mantém somente um ponto para separar milhar e a vírgula antes dos centavos.
    """
    try:
        temp = us_string.replace(",", "")
        if temp.count('.') > 1:
            parts = temp.split('.')
            temp = parts[0] + '.' + ''.join(parts[1:])
        valor_float = float(temp)
        return f"{valor_float:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return us_string

def df_to_docx_bytes(df: pd.DataFrame, titulo: str) -> bytes:
    """
    Converte um DataFrame em um arquivo DOCX (paisagem), retorna os bytes gerados.
    As linhas especiais (A, B, Indébito, Indébito em dobro) terão toda a linha em vermelho.
    """
    document = Document()
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    # Inserir Título
    head = document.add_heading(titulo, level=1)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Se existir a coluna "DESCRIÇÃO", removemos:
    if "DESCRIÇÃO" in df.columns:
        df = df.drop(columns=["DESCRIÇÃO"])

    # Linhas especiais
    linhas_especiais = [
        "A = Valor Total (R$)",
        "B = Valor Recebido - Autor (a)",
        "Indébito (A-B)",
        "Indébito em dobro (R$)"
    ]

    if df.empty:
        p = document.add_paragraph("DataFrame vazio - nenhum dado para exibir.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        colunas = df.columns.tolist()
        table = document.add_table(rows=1, cols=len(colunas))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(colunas):
            hdr_cells[i].text = str(col_name)
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Linhas do DataFrame
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            eh_linha_especial = False
            if "DISCRIMINAÇÃO" in df.columns:
                if row["DISCRIMINAÇÃO"] in linhas_especiais:
                    eh_linha_especial = True

            for i, col_name in enumerate(colunas):
                val = str(row[col_name]) if pd.notnull(row[col_name]) else ""
                paragraph = row_cells[i].paragraphs[0]
                run = paragraph.add_run(val)

                # Se for linha especial, deixar texto e valores em vermelho
                if eh_linha_especial:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                    run.font.size = Pt(14)
                    run.bold = True
                else:
                    run.font.size = Pt(9)

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ajuste de largura
        width_map = {}
        for col in colunas:
            # Tabela genérica: esticamos a "DISCRIMINAÇÃO"
            if col.upper() == "DISCRIMINAÇÃO":
                width_map[col] = 75
            else:
                width_map[col] = 25
        for i, col in enumerate(colunas):
            mm = width_map.get(col, 25)
            table.columns[i].width = Inches(mm / 25.4)

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()

def ajustar_valores_docx(file_input_bytes: bytes) -> bytes:
    """
    Recebe bytes de um arquivo docx, faz a correção de valores no texto,
    do formato '123,456.78' (US) para '123.456,78' (BR).
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_in:
        tmp_in.write(file_input_bytes)
        tmp_in.flush()
        input_path = tmp_in.name

    output_path = input_path.replace(".docx", "_corrigido.docx")
    doc = Document(input_path)
    pattern = re.compile(r'([\d,]+\.\d{2})')

    for para in doc.paragraphs:
        found = pattern.findall(para.text)
        if found:
            for val_us in found:
                val_br = formatar_valor_brl(val_us)
                para.text = para.text.replace(val_us, val_br)

    # Também corrigir tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                found = pattern.findall(txt)
                if found:
                    for val_us in found:
                        val_br = formatar_valor_brl(val_us)
                        cell.text = cell.text.replace(val_us, val_br)

    doc.save(output_path)

    with open(output_path, "rb") as f:
        final_bytes = f.read()

    os.remove(input_path)
    os.remove(output_path)
    return final_bytes

def ajustar_datas(df):
    """
    Transforma colunas JAN..DEZ + ANO em linhas do tipo:
      DATAS (Ex.: "JAN/2021"), DISCRIMINAÇÃO, VALOR (R$).
    """
    meses = ["JAN", "FEV", "MAR", "ABR", "MAI", "JUN", "JUL", "AGO", "SET", "OUT", "NOV", "DEZ"]
    linhas_ajustadas = []
    for idx, row in df.iterrows():
        ano = row.get("ANO", "")
        discriminacao = row.get("DISCRIMINAÇÃO", "")
        for mes in meses:
            valor = row.get(mes, None)
            if pd.notnull(valor):
                try:
                    valor_float = float(str(valor).replace(",", "."))  # Ajusta para float
                except:
                    valor_float = 0
                if valor_float != 0:
                    data = f"{mes}/{ano}" if str(ano).isdigit() and len(str(ano)) == 4 else mes
                    linhas_ajustadas.append({
                        "DATAS": data,
                        "DISCRIMINAÇÃO": discriminacao,
                        "VALOR (R$)": valor_float
                    })
    df_ajustado = pd.DataFrame(linhas_ajustadas, columns=["DATAS", "DISCRIMINAÇÃO", "VALOR (R$)"])
    return df_ajustado

###############################################################################
# FUNÇÃO PARA INSERIR AS 4 LINHAS (A, B, Indébito, e Indébito em dobro)
###############################################################################
def inserir_totais_na_coluna(df, col_valor):
    """
    Insere 4 linhas no final do DF:
      A = Valor Total (R$)
      B = Valor Recebido
      Indébito (A-B)
      Indébito em dobro (R$)

    O valor de B é obtido de st.session_state["valor_recebido"] (ou '0').
    O texto é injetado em "DISCRIMINAÇÃO".
    """
    if "DESCRIÇÃO" in df.columns:
        df = df.drop(columns=["DESCRIÇÃO"])

    if col_valor not in df.columns:
        return df

    def _to_float(x):
        try:
            return float(str(x).replace(',', '.').strip())
        except:
            return 0.0

    # Soma (A)
    soma = df[col_valor].apply(_to_float).sum()
    df_novo = df.copy()

    # Recupera o valor B do estado
    valor_b_str = get_state_value("valor_recebido") or "0"
    try:
        valor_b_num = float(str(valor_b_str).replace(',', '.').strip())
    except:
        valor_b_num = 0.0

    # Calcula indebito e indebito em dobro
    indebito = soma - valor_b_num
    indebito_dobro = 2 * indebito

    def en_us_format(number: float) -> str:
        return f"{number:,.2f}"

    A_str = en_us_format(soma)
    B_str = valor_b_str.strip()
    indebito_str = en_us_format(indebito)
    indebito_dobro_str = en_us_format(indebito_dobro)

    # Inserir as linhas especiais na coluna "DISCRIMINAÇÃO"
    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({
            "DISCRIMINAÇÃO": ["A = Valor Total (R$)"],
            col_valor: [A_str]
        })
    ], ignore_index=True)

    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({
            "DISCRIMINAÇÃO": ["B = Valor Recebido - Autor (a)"],
            col_valor: [B_str]
        })
    ], ignore_index=True)

    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({
            "DISCRIMINAÇÃO": ["Indébito (A-B)"],
            col_valor: [indebito_str]
        })
    ], ignore_index=True)

    df_novo = pd.concat([
        df_novo,
        pd.DataFrame({
            "DISCRIMINAÇÃO": ["Indébito em dobro (R$)"],
            col_valor: [indebito_dobro_str]
        })
    ], ignore_index=True)

    # Se existirem colunas "DATA" ou "COD", esvaziamos-nas nessas linhas especiais
    linhas_especiais = [
        "A = Valor Total (R$)",
        "B = Valor Recebido - Autor (a)",
        "Indébito (A-B)",
        "Indébito em dobro (R$)"
    ]
    mask_especial = df_novo["DISCRIMINAÇÃO"].isin(linhas_especiais)
    if "DATA" in df_novo.columns:
        df_novo.loc[mask_especial, "DATA"] = ""
    if "COD" in df_novo.columns:
        df_novo.loc[mask_especial, "COD"] = ""

    return df_novo

###############################################################################
# APLICAÇÃO STREAMLIT – FLUXO COMPLETO
###############################################################################
def main():
    # 1) Exibição da logomarca (caso o arquivo exista)
    LOGO_PATH = "MP.png"
    if os.path.exists(LOGO_PATH):
        with open(LOGO_PATH, "rb") as fimg:
            logo_b64 = base64.b64encode(fimg.read()).decode()
        st.markdown(f"""
            <div style="text-align: center;">
                <img src="data:image/png;base64,{logo_b64}" style="width:300px; height:auto;" alt="Logomarca">
            </div>
        """, unsafe_allow_html=True)

    # 2) Nome do aplicativo
    st.title("Ficha Financeira Federal")

    # Upload do PDF
    pdf_enviado = st.file_uploader("Selecione o PDF", type=["pdf"])
    if pdf_enviado is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_enviado.read())
            tmp.flush()
            caminho_pdf = tmp.name

        # (A) Extrair nome e matrícula (primeira página)
        nome, matricula = extrair_nome_e_matricula(caminho_pdf)
        set_state_value("nome_cliente", nome)
        set_state_value("matricula", matricula)

        # (B) Extração aprimorada do nome do cliente (via PyMuPDF)
        nome_cliente_extraido = extrair_nome_cliente(caminho_pdf)
        st.write("Nome do cliente extraído:", nome_cliente_extraido)
        set_state_value("nome_servidor", nome_cliente_extraido)

        # 1) DataFrame de ANO REFERÊNCIA (PÁGINA, ANO)
        st.markdown("### 1) DataFrame de ANO REFERÊNCIA (PÁGINA, ANO)")
        df_ano_celulas = extrair_celulas_interesse(caminho_pdf)
        if df_ano_celulas is not None and not df_ano_celulas.empty:
            df_ano_celulas["ANO"] = df_ano_celulas["CONTEÚDO"].apply(extrair_ultimos_quatro_digitos)
            df_ano_celulas = df_ano_celulas[["PÁGINA", "ANO"]].drop_duplicates(subset="PÁGINA")
            st.dataframe(df_ano_celulas)
            dict_anos = dict(zip(df_ano_celulas["PÁGINA"], df_ano_celulas["ANO"]))
        else:
            dict_anos = {}
            st.warning("Não foram encontradas células com ANO REFERÊNCIA (pode não existir).")

        # 2) DataFrame Consolidado (com TODAS as colunas + ANO)
        st.markdown("### 2) DataFrame Consolidado (com TODAS as colunas + ANO)")
        df_consolidado = extrair_tabelas(caminho_pdf, dict_anos)

        # Exclui o arquivo temporário (PDF) para limpar
        os.unlink(caminho_pdf)

        if df_consolidado is not None and not df_consolidado.empty:
            df_consolidado = classificar_registros_ffill(df_consolidado)
            df_consolidado["PÁGINA"] = df_consolidado["PÁGINA"].astype(int)
            # Caso alguma página não esteja no dicionário, mantemos a 'ANO' já extraída
            df_consolidado["ANO"] = df_consolidado["PÁGINA"].map(dict_anos).fillna(df_consolidado["ANO"])
            st.dataframe(df_consolidado)

            # Botão de Download em PDF (DataFrame Consolidado)
            pdf_out = os.path.join(
                tempfile.gettempdir(),
                f"extrato_financeiro_unico_{sanitizar_para_arquivo(nome_cliente_extraido)}.pdf"
            )
            salvar_em_pdf(df_consolidado, pdf_out)
            with open(pdf_out, "rb") as fpdf:
                st.download_button(
                    label="Baixar PDF (DataFrame Consolidado)",
                    data=fpdf.read(),
                    file_name=f"extrato_financeiro_unico_{sanitizar_para_arquivo(nome_cliente_extraido)}.pdf",
                    mime="application/pdf"
                )

            # 3) Análise de Descontos
            st.markdown("### 3) Análise de Descontos")
            if st.button("3.1) Filtrar Operações de Descontos"):
                df_filtrado = df_consolidado[
                    (df_consolidado["TIPO"].str.upper() == "DESCONTOS") |
                    (df_consolidado["TIPO"].str.upper() == "TIPO")
                ].copy()
                set_state_value("df_filtrado_descontos", df_filtrado)

            df_filtrado_descontos = get_state_value("df_filtrado_descontos")
            if df_filtrado_descontos is not None and not df_filtrado_descontos.empty:
                st.markdown("### 3.2) DataFrame Filtrado (Somente DESCONTOS e cabeçalho TIPO)")
                st.dataframe(df_filtrado_descontos, use_container_width=True)

                st.markdown("### 3.3) Lista das Rubricas")
                rubricas = carregar_glossario("Rubricas.txt")
                if rubricas:
                    st.dataframe(pd.DataFrame({"Rubricas": rubricas}))
                else:
                    st.warning("Glossário não encontrado ou vazio.")

                # 4) Filtrar Descontos no Glossário (Precisão Ajustável)
                st.markdown("### 4) Filtrar Descontos no Glossário (Precisão Ajustável)")
                st.write(" ")
                thresh = st.slider("Nível de Similaridade (0.1 a 1.0)", 0.1, 1.0, 0.85, 0.1)
                if st.button("Filtro com Rubricas"):
                    if not rubricas:
                        st.warning("Glossário vazio. Impossível filtrar.")
                    else:
                        threshold_value = int(thresh * 100)
                        df_somente_descontos = df_filtrado_descontos[
                            df_filtrado_descontos["TIPO"].str.upper() == "DESCONTOS"
                        ].copy()
                        df_gloss = cruzar_descontos_com_rubricas(df_somente_descontos, rubricas, threshold_value)
                        set_state_value("df_gloss", df_gloss)

                df_gloss = get_state_value("df_gloss")
                if df_gloss is not None and not df_gloss.empty:
                    st.markdown("#### 4.1) Descontos x Glossário")
                    st.dataframe(df_gloss, use_container_width=True)

                    # 5) Lista Única de Descontos
                    st.markdown("### 5) Lista Única de Descontos")
                    st.markdown("#### 5.1) Marque os itens que deseja incluir:")
                    desc_unicas = sorted(df_gloss["DISCRIMINAÇÃO"].unique())
                    selecionados = []
                    for i, desc in enumerate(desc_unicas):
                        qtd = df_gloss[df_gloss["DISCRIMINAÇÃO"] == desc].shape[0]
                        if st.checkbox(f"{desc} ({qtd}x)", key=f"chk_desc_{i}"):
                            selecionados.append(desc)

                    if st.button("Confirmar Inclusões"):
                        if not selecionados:
                            st.warning("Nenhuma descrição selecionada.")
                        else:
                            df_incl = df_gloss[df_gloss["DISCRIMINAÇÃO"].isin(selecionados)].copy()
                            set_state_value("df_incluido", df_incl)
                            st.success("Descontos selecionados com sucesso!")

                    st.markdown("#### 5.2) Lista Restante após Inclusões")
                    df_incluido = get_state_value("df_incluido")
                    if df_incluido is not None and not df_incluido.empty:
                        st.dataframe(df_incluido, use_container_width=True)

                        # 5.3) Dataframe de Datas Ajustadas
                        st.markdown("#### 5.3) Dataframe de Datas Ajustadas")
                        df_datas_ajustadas = ajustar_datas(df_incluido)
                        st.dataframe(df_datas_ajustadas, use_container_width=True)

                        # 6) Relatório Final de Descontos
                        st.markdown("### 6) Apresentar Rúbricas para Débitos (Descontos Finais)")

                        # Renomear "VALOR (R$)" para "DESCONTOS"
                        df_final = df_datas_ajustadas.copy().rename(columns={"VALOR (R$)": "DESCONTOS"})

                        # Exibe prévia em formato brasileiro na coluna 'DESCONTOS'
                        st.write("**Prévia (coluna 'DESCONTOS'):**")
                        df_preview = df_final.copy()
                        df_preview["DESCONTOS"] = df_preview["DESCONTOS"].apply(
                            lambda x: "R$ " + formatar_valor_brl(str(x))
                        )
                        st.dataframe(df_preview)

                        # Soma (A), manipulação de B, etc.
                        def _to_float(x):
                            try:
                                return float(str(x).replace(',', '.').strip())
                            except:
                                return 0.0

                        A_val = df_final["DESCONTOS"].apply(_to_float).sum()
                        A_str = f"{A_val:,.2f}"

                        col1, col2 = st.columns(2)
                        with col1:
                            valor_b_receb = st.text_input("B = Valor Recebido - Autor (a) [utilizar ponto para separar os centavos]", "0")
                        try:
                            vrnum = float(valor_b_receb.replace(',', '.').strip())
                        except:
                            vrnum = 0.0

                        indebito = A_val - vrnum
                        indebito_dobro = 2 * indebito
                        indebito_str = f"{indebito:,.2f}"
                        indebito_dobro_str = f"{indebito_dobro:,.2f}"

                        with col2:
                            st.write(f"Indébito (A-B): {indebito_str}")
                            st.write(f"Indébito em dobro (R$): {indebito_dobro_str}")

                        # Armazena "valor_recebido" no estado
                        set_state_value("valor_recebido", valor_b_receb)

                        # Botão para gerar relatório final
                        with st.form("form_descontos_finais"):
                            submit_final = st.form_submit_button("Gerar Relatório Final de Descontos")

                        if submit_final:
                            titulo_final = "Descontos Finais"
                            # Monta DataFrame com as 4 linhas especiais
                            df_com_totais = inserir_totais_na_coluna(df_final.copy(), "DESCONTOS")

                            # =============== GERAÇÃO DO PDF FINAL ===============
                            from fpdf import FPDF

                            class PDFDescontosFinais(FPDF):
                                def header(self):
                                    self.set_font("Arial", "B", 16)
                                    self.cell(0, 10, titulo_final, border=False, ln=True, align='C')
                                    self.ln(5)

                                def footer(self):
                                    self.set_y(-15)
                                    self.set_font('Arial', 'I', 8)
                                    self.cell(0, 10, f'Página {self.page_no()}', border=False, ln=False, align='C')

                            pdf_temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
                            pdf_doc = PDFDescontosFinais(orientation="L", format="A4")
                            pdf_doc.add_page()

                            # Remover "DESCRIÇÃO" se ainda existir
                            if "DESCRIÇÃO" in df_com_totais.columns:
                                df_com_totais = df_com_totais.drop(columns=["DESCRIÇÃO"])

                            colunas_final = df_com_totais.columns.tolist()
                            col_widths = []
                            for c in colunas_final:
                                if c.upper() == "DISCRIMINAÇÃO":
                                    col_widths.append(150)
                                elif c.upper() == "DATAS":
                                    col_widths.append(40)
                                else:
                                    col_widths.append(40)

                            # Cabeçalho
                            pdf_doc.set_font("Arial", "B", 10)
                            for i, col in enumerate(colunas_final):
                                pdf_doc.cell(col_widths[i], 8, col, border=1, align='C')
                            pdf_doc.ln()

                            # Linhas especiais para destacar
                            linhas_especiais = [
                                "A = Valor Total (R$)",
                                "B = Valor Recebido - Autor (a)",
                                "Indébito (A-B)",
                                "Indébito em dobro (R$)"
                            ]

                            # Impressão das linhas do PDF
                            for _, row_ in df_com_totais.iterrows():
                                is_special_line = False
                                if row_["DISCRIMINAÇÃO"] in linhas_especiais:
                                    is_special_line = True

                                for i, col in enumerate(colunas_final):
                                    val = str(row_[col]) if pd.notnull(row_[col]) else ""

                                    # Para a linha "B = Valor Recebido - Autor (a)" no PDF,
                                    # dividir valor inserido pelo usuário por 10.
                                    if (row_["DISCRIMINAÇÃO"] == "B = Valor Recebido - Autor (a)") and (col.upper() == "DESCONTOS"):
                                        # Tentar converter e dividir por 10
                                        try:
                                            val_float = float(val.replace(',', '.').strip())

                                            val = f"{val_float:.2f}"
                                        except:
                                            pass

                                    # Se for coluna de valores (DESCONTOS), converter para BR
                                    if col.upper() == "DESCONTOS" and not (val in linhas_especiais):
                                        val = formatar_valor_brl(val)

                                    # Configurar cores/fonte se for linha especial
                                    if is_special_line:
                                        pdf_doc.set_text_color(255, 0, 0)     # Vermelho
                                        pdf_doc.set_font("Arial", "B", 12)    # Negrito, maior
                                    else:
                                        pdf_doc.set_text_color(0, 0, 0)
                                        pdf_doc.set_font("Arial", "", 10)

                                    pdf_doc.cell(col_widths[i], 8, val, border=1, align='C')
                                pdf_doc.ln()

                            pdf_doc.output(pdf_temp)
                            with open(pdf_temp, "rb") as fpdf_:
                                pdf_data_finais = fpdf_.read()
                            os.remove(pdf_temp)

                            pdf_download_name = f"Descontos_Finais_Cronologico_{sanitizar_para_arquivo(nome_cliente_extraido)}.pdf"
                            st.download_button(
                                label="Baixar PDF (Descontos Finais - Cronológico)",
                                data=pdf_data_finais,
                                file_name=pdf_download_name,
                                mime="application/pdf"
                            )

                            # =============== GERAÇÃO DO DOCX FINAL ===============
                            docx_data = df_to_docx_bytes(df_com_totais, titulo_final)
                            # Ajustar valores para formato BR (via regex)
                            docx_corrigido = ajustar_valores_docx(docx_data)
                            docx_download_name = f"Descontos_Finais_Cronologico_{sanitizar_para_arquivo(nome_cliente_extraido)}.docx"

                            st.download_button(
                                label="Baixar DOCX (Descontos Finais - Cronológico)",
                                data=docx_corrigido,
                                file_name=docx_download_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

if __name__ == "__main__":
    st.set_page_config(page_title="Analista de Contracheques", layout="centered")
    main()
